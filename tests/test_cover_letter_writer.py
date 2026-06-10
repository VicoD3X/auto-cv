from docx import Document

from autocv.documents import CoverLetterDocxWriter, CoverLetterWriteRequest


def test_cover_letter_writer_replaces_body_and_preserves_template(tmp_path) -> None:
    template = tmp_path / "template.docx"
    output = tmp_path / "Result" / "letter.docx"
    _create_template(template)

    writer = CoverLetterDocxWriter()
    writer.write(
        CoverLetterWriteRequest(
            template_path=template,
            output_path=output,
            body_text="Nouveau paragraphe un.\n\nNouveau paragraphe deux.",
            project_name="spark-vision",
            project_url="https://github.com/VicoD3X/spark-vision",
        )
    )

    source_text = "\n".join(paragraph.text for paragraph in Document(str(template)).paragraphs)
    output_doc = Document(str(output))
    output_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
    hyperlink_targets = [
        rel.target_ref for rel in output_doc.part.rels.values() if "hyperlink" in rel.reltype
    ]

    assert "Ancien corps" in source_text
    assert "Ancien corps" not in output_text
    assert "Nouveau paragraphe un." in output_text
    assert "spark-vision" in output_text
    assert hyperlink_targets == ["https://github.com/VicoD3X/spark-vision"]


def _create_template(path) -> None:
    document = Document()
    for text in [
        "Victor Aubry",
        "Metz | github.com/VicoD3X",
        "À l’attention du service recrutement de",
        "Metz, le 11 mai 2026",
        "Objet : Candidature au poste de Data scientist",
        "Madame, Monsieur,",
        "Ancien corps un.",
        "Ancien corps deux.",
        "Je vous prie d’agréer, Madame, Monsieur, l’expression de mes salutations distinguées.",
        "Victor Aubry",
    ]:
        document.add_paragraph(text)
    document.save(str(path))
