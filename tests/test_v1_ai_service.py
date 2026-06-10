from pathlib import Path

from docx import Document

from autocv.documents.naming import DocumentKind
from autocv.domain import ApplicationRecord, OpportunityType
from autocv.engine import EngineRequest, EngineResponse
from autocv.mail import MailDraftRequest
from autocv.projects import GitHubProjectContext
from autocv.use_cases import V1AiService


class FakeEngine:
    def generate(self, request: EngineRequest) -> EngineResponse:
        if request.task == "mail_draft":
            return EngineResponse(
                text="Objet: Candidature Data Scientist\nCorps: Bonjour,\nVoici ma candidature.",
                source="fake",
                available=True,
                model="fake-model",
            )
        return EngineResponse(
            text=f"generated:{request.task}:{request.context['target_name']}",
            source="fake",
            available=True,
            model="fake-model",
        )


def test_v1_ai_service_routes_job_letter_adaptation() -> None:
    service = V1AiService(engine=FakeEngine())
    record = _record(OpportunityType.JOB)

    response = service.adapt_application_text(
        record=record,
        target_name="Airbus",
        role_or_mission="Data Scientist",
        context="Offre Python ML",
    )

    assert response.available is True
    assert response.text == "generated:cover_letter_adaptation:Airbus"


def test_v1_ai_service_routes_free_chat_with_project_context() -> None:
    service = V1AiService(engine=FakeEngine())
    project = GitHubProjectContext(
        repository_name="spark-vision",
        url="https://github.com/VicoD3X/spark-vision",
        description="Vision pipeline",
    )

    response = service.chat(
        message="Pourquoi ce projet colle ?",
        record=_record(OpportunityType.JOB),
        target_name="Airbus",
        role_or_mission="Data Scientist",
        context="Offre ML",
        selected_project=project,
        chat_history="user: hello",
    )

    assert response.available is True
    assert response.text == "generated:chat:Airbus"


def test_v1_ai_service_routes_freelance_proposal() -> None:
    service = V1AiService(engine=FakeEngine())
    record = _record(OpportunityType.FREELANCE)

    response = service.adapt_application_text(
        record=record,
        target_name="Client",
        role_or_mission="Dashboard",
        context="Besoin KPI",
    )

    assert response.text == "generated:freelance_proposal:Client"


def test_v1_ai_service_parses_mail_draft() -> None:
    service = V1AiService(engine=FakeEngine())

    draft = service.draft_mail(
        request=MailDraftRequest(
            opportunity_type=OpportunityType.JOB,
            target_name="Airbus",
            role_or_mission="Data Scientist",
            context="Candidature",
        )
    )

    assert draft.available is True
    assert draft.subject == "Candidature Data Scientist"
    assert draft.body == "Bonjour,\nVoici ma candidature."


def test_v1_ai_service_saves_preview_to_result_dir(tmp_path) -> None:
    service = V1AiService(engine=FakeEngine())

    output = service.save_preview(
        result_dir=tmp_path,
        kind=DocumentKind.EMAIL_DRAFT,
        target_name="Airbus",
        role_or_mission="Data Scientist",
        date="2026-06-09",
        content="Bonjour",
    )

    assert output == Path(tmp_path) / "Mail_Airbus_Data_Scientist_2026_06_09.txt"
    assert output.read_text(encoding="utf-8") == "Bonjour"


def test_v1_ai_service_generates_cover_letter_docx(tmp_path) -> None:
    template = tmp_path / "lettre.docx"
    _create_letter_template(template)
    record = ApplicationRecord.create(
        opportunity_type=OpportunityType.JOB,
        opportunity_id="opportunity-1",
        cv_path="cv.pdf",
        cover_letter_source_path=str(template),
        cover_letter_output_path=str(tmp_path / "Result" / "lettre-output.docx"),
    )
    project = GitHubProjectContext(
        repository_name="spark-vision",
        url="https://github.com/VicoD3X/spark-vision",
    )
    service = V1AiService(engine=FakeEngine())

    result = service.generate_cover_letter_docx(
        record=record,
        target_name="Airbus",
        role_or_mission="Data Scientist",
        context="Offre ML",
        result_dir=tmp_path / "Result",
        selected_project=project,
    )

    output_text = "\n".join(paragraph.text for paragraph in Document(str(result.output_path)).paragraphs)

    assert result.available is True
    assert result.output_path == tmp_path / "Result" / "lettre-output.docx"
    assert "generated:cover_letter_adaptation:Airbus" in output_text
    assert "spark-vision" in output_text
    assert "Ancien corps" not in output_text


def _record(opportunity_type: OpportunityType) -> ApplicationRecord:
    return ApplicationRecord.create(
        opportunity_type=opportunity_type,
        opportunity_id="opportunity-1",
        cv_path="cv.pdf",
        cover_letter_source_path="lettre.docx",
    )


def _create_letter_template(path: Path) -> None:
    document = Document()
    for text in [
        "Victor Aubry",
        "Metz | github.com/VicoD3X",
        "À l’attention du service recrutement de",
        "Metz, le 11 mai 2026",
        "Objet : Candidature au poste de Data scientist",
        "Madame, Monsieur,",
        "Ancien corps.",
        "Je vous prie d’agréer, Madame, Monsieur, l’expression de mes salutations distinguées.",
        "Victor Aubry",
    ]:
        document.add_paragraph(text)
    document.save(str(path))
