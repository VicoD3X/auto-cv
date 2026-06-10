from dataclasses import dataclass
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


@dataclass(frozen=True, slots=True)
class CoverLetterWriteRequest:
    template_path: Path
    output_path: Path
    body_text: str
    project_name: str = ""
    project_url: str = ""


class CoverLetterDocxWriter:
    def write(self, request: CoverLetterWriteRequest) -> Path:
        if not request.template_path.exists():
            raise FileNotFoundError(request.template_path)

        request.output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(request.template_path, request.output_path)

        document = Document(str(request.output_path))
        body_start, body_end = _body_bounds(document)
        body_style = document.paragraphs[body_start].style if body_start < len(document.paragraphs) else None
        anchor = document.paragraphs[body_end] if body_end < len(document.paragraphs) else None

        for paragraph in list(document.paragraphs[body_start:body_end]):
            _remove_paragraph(paragraph)

        inserted = []
        for text in _split_body(request.body_text):
            if anchor is None:
                inserted.append(document.add_paragraph(text, style=body_style))
            else:
                inserted.append(anchor.insert_paragraph_before(text, style=body_style))

        if request.project_name and request.project_url:
            target = inserted[-1] if inserted else (anchor or document.add_paragraph(style=body_style))
            if target.text:
                target.add_run(" ")
            target.add_run("Projet GitHub: ")
            _add_hyperlink(target, request.project_name, request.project_url)

        document.save(str(request.output_path))
        return request.output_path


def _body_bounds(document: Document) -> tuple[int, int]:
    paragraphs = document.paragraphs
    if not paragraphs:
        return 0, 0

    body_start = 6 if len(paragraphs) > 6 else 0
    for index, paragraph in enumerate(paragraphs):
        normalized = paragraph.text.strip().lower()
        if "madame" in normalized and "monsieur" in normalized:
            body_start = min(index + 1, len(paragraphs))
            break

    body_end = max(body_start, len(paragraphs) - 2)
    for index in range(body_start, len(paragraphs)):
        normalized = paragraphs[index].text.strip().lower()
        if "je vous prie" in normalized or "salutations distingu" in normalized:
            body_end = index
            break

    return body_start, max(body_start, body_end)


def _split_body(text: str) -> list[str]:
    parts = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n")]
    paragraphs: list[str] = []
    for part in parts:
        lines = [line.strip() for line in part.splitlines() if line.strip()]
        if lines:
            paragraphs.append(" ".join(lines))
    if not paragraphs and text.strip():
        paragraphs.append(text.strip())
    return paragraphs


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
